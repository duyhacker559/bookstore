#!/usr/bin/env python3
"""
Generate Technical Report DOCX (Word) for Bookstore Microservices System
Editable version for easy modification and updates
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path

def add_heading_style(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(31, 71, 136)
        heading.runs[0].font.size = Pt(24) if level == 1 else Pt(16)
    return heading

def add_table_with_style(doc, data, header=True):
    """Add a styled table"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Light Grid Accent 1'
    
    if header:
        # Style header row
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '2c5aa0')
            cell._element.get_or_add_tcPr().append(shading_elm)
            # Make header text white and bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
    
    # Fill table data
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = str(cell_data)
    
    return table

def shade_paragraph(paragraph, color='3a6bb0'):
    """Add background shading to paragraph"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    paragraph._element.get_or_add_tcPr().append(shading_elm)


def get_roadmap_diagram_paths():
    """Return sorted rendered roadmap diagram paths for Section 18."""
    render_dir = Path(r'l:\bookstore\diagrams\rendered')
    patterns = [
        'roadmap-*.png',
        'roadmap-*.jpg',
        'roadmap-*.jpeg',
    ]

    diagram_paths = []
    for pattern in patterns:
        diagram_paths.extend(render_dir.glob(pattern))

    return sorted(set(diagram_paths), key=lambda p: p.name.lower())

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ===== TITLE PAGE =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('TECHNICAL REPORT')
title_run.font.size = Pt(28)
title_run.font.color.rgb = RGBColor(31, 71, 136)
title_run.font.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Bookstore Microservices Architecture')
subtitle_run.font.size = Pt(18)
subtitle_run.font.color.rgb = RGBColor(31, 71, 136)
subtitle_run.font.bold = True

doc.add_paragraph()

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2_run = subtitle2.add_run('Event-Driven Implementation with Django, FastAPI, and RabbitMQ')
subtitle2_run.font.size = Pt(12)
subtitle2_run.font.italic = True
subtitle2_run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
doc.add_paragraph()

# Document info
info_table_data = [
    ['Date:', f'{datetime.now().strftime("%B %d, %Y")}'],
    ['System Type:', 'Hybrid Monolith + Microservices'],
    ['Framework:', 'Django + FastAPI'],
    ['Message Broker:', 'RabbitMQ'],
    ['Container Orchestration:', 'Docker Compose'],
    ['Primary Language:', 'Python'],
]
info_table = add_table_with_style(doc, info_table_data, header=False)

# Page break
doc.add_page_break()

# ===== ABSTRACT =====
add_heading_style(doc, 'Abstract', level=1)
abstract_text = (
    "This report presents a production-ready microservices architecture for an e-commerce bookstore "
    "system utilizing a hybrid approach: a monolithic Django application for web UI and orchestration, "
    "combined with three independent FastAPI microservices for Payment, Shipping, and Notification processing. "
    "The system implements event-driven communication via RabbitMQ for asynchronous workflows, synchronous "
    "REST APIs for critical-path operations, and database-per-service isolation with PostgreSQL. "
    "The architecture achieves separation of concerns, independent scaling, graceful fault handling, and "
    "clear service boundaries following Domain-Driven Design principles. This report details the architectural "
    "design, service responsibilities, communication patterns, deployment model, and provides implementation "
    "walkthroughs for key business scenarios."
)
doc.add_paragraph(abstract_text)

keywords = doc.add_paragraph()
keywords_run = keywords.add_run("Keywords: ")
keywords_run.font.bold = True
keywords.add_run(
    "microservices, Django REST Framework, FastAPI, RabbitMQ, event-driven architecture, "
    "API gateway design, database-per-service, Docker Compose, domain-driven design, e-commerce."
)

doc.add_page_break()

# ===== 1. INTRODUCTION =====
add_heading_style(doc, '1. Introduction', level=1)

add_heading_style(doc, '1.1 Background', level=2)
doc.add_paragraph(
    "Traditional e-commerce platforms face scaling challenges as business complexity grows. "
    "A single application instance managing user authentication, catalog browsing, shopping carts, payment processing, "
    "and shipment tracking creates resource contention and deployment risk. The Bookstore system addresses this by "
    "decomposing operational concerns into independent services while maintaining a centralized web interface through "
    "a Django monolith that orchestrates customer-facing workflows."
)

add_heading_style(doc, '1.2 Problem Statement', level=2)
doc.add_paragraph(
    "A purely monolithic bookstore architecture presents several operational constraints:"
)

problems = [
    "Uneven resource utilization: Payment and shipping operations require different computational profiles and availability SLOs than catalog browsing.",
    "Deployment coupling: Each update to payment logic forces redeployment of the entire system.",
    "Technology constraints: A single technology stack may not suit specialized workloads (e.g., real-time notifications).",
    "Scaling granularity: Scaling cannot target specific bottlenecks independently.",
    "Team organization: Tight coupling complicates parallel development across functional domains."
]

for problem in problems:
    p = doc.add_paragraph(problem, style='List Bullet')

add_heading_style(doc, '1.3 Objective', level=2)
doc.add_paragraph(
    "This project implements a hybrid microservices architecture that balances pragmatism with architectural "
    "best practices. The solution maintains a monolithic web tier for customer UI and order orchestration while "
    "delegating payment, shipping, and notification concerns to dedicated FastAPI services. Event-driven communication "
    "decouples the notification service from transactional workloads, enabling eventual consistency and resilience."
)

doc.add_page_break()

# ===== 2. ARCHITECTURAL OVERVIEW =====
add_heading_style(doc, '2. Architectural Overview', level=1)

add_heading_style(doc, '2.1 System Topology', level=2)
doc.add_paragraph("The system employs a hybrid architecture with the following components:")

topology_data = [
    ['Component', 'Type', 'Purpose', 'Port'],
    ['Django Web App', 'Monolith', 'UI, authentication, cart, order orchestration', '8000'],
    ['Payment Service', 'FastAPI', 'Payment processing and refunds', '5000'],
    ['Shipping Service', 'FastAPI', 'Shipment creation and tracking', '5001'],
    ['Notification Service', 'FastAPI', 'Email notifications (event consumer)', '5002'],
    ['PostgreSQL', 'Database', '4 databases (bookstore, payment, shipping, notification)', '5432'],
    ['RabbitMQ', 'Message Broker', 'Asynchronous event distribution', '5672/15672'],
]

add_table_with_style(doc, topology_data, header=True)

add_heading_style(doc, '2.2 Communication Patterns', level=2)
doc.add_paragraph("The system implements two primary communication patterns:")

sync_p = doc.add_paragraph(style='List Bullet')
sync_run = sync_p.add_run("Synchronous (HTTP/REST): ")
sync_run.bold = True
sync_p.add_run(
    "Used for critical-path operations requiring immediate responses "
    "(payment processing and shipment creation). Requests include bearer token authentication and implement "
    "timeouts for graceful degradation."
)

async_p = doc.add_paragraph(style='List Bullet')
async_run = async_p.add_run("Asynchronous (RabbitMQ): ")
async_run.bold = True
async_p.add_run(
    "Used for notification workflows where loose coupling and eventual consistency "
    "are acceptable. The Notification Service subscribes to domain events published by Payment and Shipping services, "
    "enabling independent deployment and fault isolation."
)

add_heading_style(doc, '2.3 Database Architecture', level=2)
db_text = doc.add_paragraph(
    "Each service owns independent PostgreSQL databases, preventing cross-service schema coupling:"
)

db_data = [
    ['Database', 'Service', 'Key Tables', 'Purpose'],
    ['bookstore', 'Django Monolith', 'users, orders, books, cart, customers', 'Web tier and order orchestration'],
    ['payment_service', 'Payment Service', 'payments, refunds', 'Payment transaction records'],
    ['shipping_service', 'Shipping Service', 'shipments', 'Shipment tracking'],
    ['notification_service', 'Notification Service', 'notifications', 'Notification audit trail'],
]

add_table_with_style(doc, db_data, header=True)

doc.add_page_break()

# ===== 3. SERVICE DESIGN =====
add_heading_style(doc, '3. Service Design and Responsibilities', level=1)

add_heading_style(doc, '3.1 Django Monolith', level=2)
doc.add_paragraph("The Django application serves as the primary customer interface and orchestrator:")

resp_heading = doc.add_paragraph()
resp_run = resp_heading.add_run("Core Responsibilities:")
resp_run.bold = True

responsibilities = [
    "User authentication (username/password with token generation)",
    "Book catalog management and search",
    "Shopping cart operations (add, view, update quantities)",
    "Order creation and status tracking",
    "Integration with external payment service (synchronous calls)",
    "Integration with external shipping service (synchronous calls)",
    "Event publishing for order creation and updates",
    "Staff management UI for shipment status updates"
]

for resp in responsibilities:
    doc.add_paragraph(resp, style='List Bullet')

model_heading = doc.add_paragraph()
model_run = model_heading.add_run("Key Models:")
model_run.bold = True

models = [
    "User: username, email, password, authentication token",
    "Customer: user reference, loyalty points, registration date",
    "Book: ISBN, title, author, price, stock quantity",
    "Order: customer reference, total amount, status (pending/processing/completed/cancelled)",
    "CartItem: cart reference, book reference, quantity"
]

for model in models:
    doc.add_paragraph(model, style='List Bullet')

add_heading_style(doc, '3.2 Payment Service', level=2)
service_info = doc.add_paragraph(
    "Framework: FastAPI 0.104.1 | Port: 5000 | Database: payment_service (PostgreSQL)"
)
service_info.runs[0].italic = True

resp_heading = doc.add_paragraph()
resp_run = resp_heading.add_run("Core Responsibilities:")
resp_run.bold = True

payment_resp = [
    "Process payment transactions via Stripe API",
    "Manage payment status lifecycle (PENDING → PROCESSING → COMPLETED/FAILED)",
    "Handle refund requests and tracking",
    "Implement idempotency checking to prevent duplicate charges",
    "Publish 'order.paid' event to message broker on successful payment",
    "Provide payment status queries for order tracking"
]

for resp in payment_resp:
    doc.add_paragraph(resp, style='List Bullet')

endpoint_heading = doc.add_paragraph()
endpoint_run = endpoint_heading.add_run("API Endpoints:")
endpoint_run.bold = True

endpoints = [
    "POST /api/v1/payments/process – Submit payment for processing",
    "GET /api/v1/payments/{order_id} – Retrieve payment status",
    "POST /api/v1/payments/refund – Request payment refund",
    "GET /health – Health check probe"
]

for ep in endpoints:
    doc.add_paragraph(ep, style='List Bullet')

add_heading_style(doc, '3.3 Shipping Service', level=2)
doc.add_paragraph(
    "Framework: FastAPI 0.104.1 | Port: 5001 | Database: shipping_service (PostgreSQL)"
).runs[0].italic = True

resp_heading = doc.add_paragraph()
resp_run = resp_heading.add_run("Core Responsibilities:")
resp_run.bold = True

shipping_resp = [
    "Calculate shipping costs based on method and destination",
    "Create shipment records and assign tracking numbers",
    "Update shipment status through lifecycle (PENDING → PROCESSING → SHIPPED → DELIVERED/FAILED)",
    "Publish domain events (shipment.created, shipment.updated) to message broker",
    "Provide shipping options to customers (Standard $5, Express $15, Overnight $50)",
    "Track estimated delivery dates"
]

for resp in shipping_resp:
    doc.add_paragraph(resp, style='List Bullet')

methods_heading = doc.add_paragraph()
methods_run = methods_heading.add_run("Shipping Methods:")
methods_run.bold = True

methods = [
    "Standard Shipping: $5 (5-7 business days)",
    "Express Shipping: $15 (2-3 business days)",
    "Overnight Shipping: $50 (next business day)"
]

for method in methods:
    doc.add_paragraph(method, style='List Bullet')

add_heading_style(doc, '3.4 Notification Service', level=2)
doc.add_paragraph(
    "Framework: FastAPI 0.104.1 | Port: 5002 | Database: notification_service (PostgreSQL)"
).runs[0].italic = True

resp_heading = doc.add_paragraph()
resp_run = resp_heading.add_run("Core Responsibilities:")
resp_run.bold = True

notif_resp = [
    "Consume domain events from RabbitMQ (asynchronous worker)",
    "Generate and send email notifications based on event types",
    "Maintain notification audit trail in database",
    "Handle email delivery failures and retries",
    "Support both development (console logging) and production (SMTP) modes",
    "Implement automatic reconnection and error recovery"
]

for resp in notif_resp:
    doc.add_paragraph(resp, style='List Bullet')

event_heading = doc.add_paragraph()
event_run = event_heading.add_run("Events Consumed:")
event_run.bold = True

events = [
    "order.paid – Sends order confirmation email to customer",
    "shipment.created – Sends shipping notification with tracking information",
    "shipment.updated – Sends status update emails for delivered/failed shipments"
]

for event in events:
    doc.add_paragraph(event, style='List Bullet')

doc.add_page_break()

# ===== 4. END-TO-END FLOW =====
add_heading_style(doc, '4. End-to-End Order Processing Flow', level=1)

add_heading_style(doc, '4.1 Complete Order Lifecycle', level=2)
doc.add_paragraph(
    "The following sequence illustrates the complete order processing workflow from customer "
    "submission through shipment:"
)

flow_steps = [
    "1. Customer submits order through Django web interface with selected books and quantities.",
    "2. Order is validated and saved in bookstore database with status=PENDING.",
    "3. Django monolith calls Payment Service synchronously via HTTP POST /api/v1/payments/process.",
    "4. Payment Service processes payment with Stripe, updates payment record, and publishes order.paid event to RabbitMQ.",
    "5. Django receives payment confirmation and calls Shipping Service synchronously via HTTP POST /api/v1/shipping/create.",
    "6. Shipping Service creates shipment record, publishes shipment.created event to RabbitMQ.",
    "7. Order status is updated to PROCESSING in Django database.",
    "8. Client receives immediate confirmation with order and tracking details.",
    "9. [Asynchronously] Notification Service consumes order.paid event, sends order confirmation email.",
    "10. [Asynchronously] Notification Service consumes shipment.created event, sends shipping notification email with tracking.",
    "11. Notification records are persisted in notification_service database for audit trail."
]

for step in flow_steps:
    doc.add_paragraph(step, style='List Number')

add_heading_style(doc, '4.2 Failure Handling', level=2)
doc.add_paragraph("The system implements graceful degradation for service outages:")

failure_p1 = doc.add_paragraph(style='List Bullet')
failure1_run = failure_p1.add_run("Payment Service Unavailable: ")
failure1_run.bold = True
failure_p1.add_run(
    "Request times out after 30 seconds; customer is notified and can retry. "
    "Transaction state is not corrupted."
)

failure_p2 = doc.add_paragraph(style='List Bullet')
failure2_run = failure_p2.add_run("Shipping Service Unavailable: ")
failure2_run.bold = True
failure_p2.add_run(
    "Order is marked for manual processing; staff can create shipment manually through UI. "
    "Customer receives partial confirmation."
)

failure_p3 = doc.add_paragraph(style='List Bullet')
failure3_run = failure_p3.add_run("Notification Service Unavailable: ")
failure3_run.bold = True
failure_p3.add_run(
    "Events remain in RabbitMQ queue and will be processed once service recovers. "
    "This does not affect order completion."
)

doc.add_page_break()

# ===== 5. EVENT-DRIVEN ARCHITECTURE =====
add_heading_style(doc, '5. Event-Driven Architecture', level=1)

add_heading_style(doc, '5.1 RabbitMQ Configuration', level=2)

config_data = [
    ['Property', 'Value', 'Purpose'],
    ['Message Broker', 'RabbitMQ 3', 'Asynchronous event distribution'],
    ['Exchange Type', 'Topic', 'Content-based routing of events'],
    ['Exchange Name', 'bookstore.events', 'Central event bus'],
    ['Queue Name', 'notification_service_queue', 'Notification service subscription'],
    ['Connection Port', '5672 (AMQP)', 'Inter-service communication'],
    ['Management UI', '15672 (HTTP)', 'Broker monitoring and admin'],
    ['Durability', 'Enabled', 'Messages survive broker restart'],
]

add_table_with_style(doc, config_data, header=True)

add_heading_style(doc, '5.2 Event Message Format', level=2)
doc.add_paragraph("All domain events follow a consistent structure for interoperability:")

code_p = doc.add_paragraph()
code_run = code_p.add_run(
    '{ "event_type": "order.paid", "timestamp": "2026-03-11T10:30:45Z", '
    '"correlation_id": "uuid-here", "data": {"order_id": 123, "amount": 99.99, ...} }'
)
code_run.font.name = 'Courier New'
code_run.font.color.rgb = RGBColor(31, 71, 136)

add_heading_style(doc, '5.3 Published Events', level=2)

pub_event1 = doc.add_paragraph(style='List Bullet')
pub1_run = pub_event1.add_run("order.paid ")
pub1_run.bold = True
pub_event1.add_run(
    "– Published by Payment Service after successful payment. Contains "
    "order_id, payment_id, amount, timestamp. Triggers order confirmation email."
)

pub_event2 = doc.add_paragraph(style='List Bullet')
pub2_run = pub_event2.add_run("shipment.created ")
pub2_run.bold = True
pub_event2.add_run(
    "– Published by Shipping Service when shipment record is created. "
    "Contains order_id, shipment_id, tracking_number, estimated_delivery. Triggers shipping notification email."
)

pub_event3 = doc.add_paragraph(style='List Bullet')
pub3_run = pub_event3.add_run("shipment.updated ")
pub3_run.bold = True
pub_event3.add_run(
    "– Published by Shipping Service on status changes (SHIPPED, DELIVERED, FAILED). "
    "Triggers customer status update emails."
)

doc.add_page_break()

# ===== 6. DEPLOYMENT =====
add_heading_style(doc, '6. Deployment and Infrastructure', level=1)

add_heading_style(doc, '6.1 Docker Compose Orchestration', level=2)
doc.add_paragraph(
    "The system is containerized and orchestrated via Docker Compose for development and testing. "
    "Each service defines its own Dockerfile and runs independently within the Docker network."
)

services_heading = doc.add_paragraph()
services_run = services_heading.add_run("Services Managed by Docker Compose:")
services_run.bold = True

services = [
    "web (Django): Port 8000, hot reload enabled for development, depends on postgres and rabbitmq",
    "payment-service (FastAPI): Port 5000, auto-reload enabled, depends on postgres and rabbitmq",
    "shipping-service (FastAPI): Port 5001, auto-reload enabled, depends on postgres and rabbitmq",
    "notification-service (FastAPI): Port 5002, background RabbitMQ consumer, depends on rabbitmq",
    "postgres: Port 5432, PostgreSQL 14 database server for all services",
    "rabbitmq: Ports 5672 (AMQP) and 15672 (management UI)"
]

for svc in services:
    doc.add_paragraph(svc, style='List Bullet')

add_heading_style(doc, '6.2 Health Checks', level=2)
doc.add_paragraph("All services implement health check endpoints that Docker Compose monitors:")

health_p = doc.add_paragraph(style='List Bullet')
health_run = health_p.add_run("GET /health ")
health_run.bold = True
health_p.add_run(
    "– Returns 200 OK if service is operational. Used by Docker to verify "
    "readiness before accepting traffic."
)

add_heading_style(doc, '6.3 Quick Start', level=2)
doc.add_paragraph("To run the complete system:")

cmd_p = doc.add_paragraph()
cmd_run = cmd_p.add_run("$ docker-compose up")
cmd_run.font.name = 'Courier New'
cmd_run.font.bold = True

doc.add_paragraph(
    "This command starts all services, applies database migrations, creates schema, and initializes RabbitMQ queues. "
    "The system is ready for testing after all health checks pass."
)

doc.add_page_break()

# ===== 7. SECURITY =====
add_heading_style(doc, '7. Security and Authentication', level=1)

add_heading_style(doc, '7.1 Service-to-Service Authentication', level=2)
doc.add_paragraph("Inter-service communication is authenticated using bearer tokens:")

doc.add_paragraph(
    "All HTTP requests from Django monolith to Payment/Shipping services include an Authorization header "
    "with a bearer token. Services validate the token before processing requests. Each service has a unique token "
    "configured in environment variables."
)

add_heading_style(doc, '7.2 Database Access Control', level=2)
doc.add_paragraph("Each service connects to PostgreSQL using dedicated credentials with principle of least privilege:")

doc.add_paragraph(
    "Services authenticate with PostgreSQL using username/password credentials over connection pooling. "
    "Network access is restricted within the Docker network (no external database access)."
)

add_heading_style(doc, '7.3 Message Broker Security', level=2)
doc.add_paragraph(
    "RabbitMQ access is controlled via username/password (default guest/guest for development). "
    "In production, strong credentials and TLS encryption would be required."
)

add_heading_style(doc, '7.4 Recommended Security Enhancements', level=2)

enhancements = [
    "Implement JWT tokens for stateless authentication across services",
    "Add rate limiting and request throttling at the API Gateway (Django)",
    "Enable TLS/SSL for inter-service HTTP communication and RabbitMQ",
    "Implement request signing to prevent tampering",
    "Add comprehensive audit logging for sensitive operations",
    "Enable encryption at rest for database storage"
]

for enh in enhancements:
    doc.add_paragraph(enh, style='List Bullet')

doc.add_page_break()

# ===== 8. SCALABILITY =====
add_heading_style(doc, '8. Scalability and Performance Considerations', level=1)

add_heading_style(doc, '8.1 Horizontal Scaling', level=2)
doc.add_paragraph("Services can be scaled independently based on demand:")

scale_p1 = doc.add_paragraph(style='List Bullet')
scale1_run = scale_p1.add_run("Catalog Browsing (Read-Heavy): ")
scale1_run.bold = True
scale_p1.add_run(
    "Scale Django and book-related queries independently. "
    "Implement caching to reduce database load."
)

scale_p2 = doc.add_paragraph(style='List Bullet')
scale2_run = scale_p2.add_run("Payment Processing (CPU/Network Intensive): ")
scale2_run.bold = True
scale_p2.add_run(
    "Scale Payment Service independently. "
    "Payment transactions block order flow, so throughput directly impacts customer experience."
)

scale_p3 = doc.add_paragraph(style='List Bullet')
scale3_run = scale_p3.add_run("Notifications (Queue-Based): ")
scale3_run.bold = True
scale_p3.add_run(
    "Scale Notification Service based on message backlog. "
    "This workload is decoupled and doesn't affect customer-facing operations."
)

add_heading_style(doc, '8.2 Performance Optimization Strategies', level=2)

optimizations = [
    "Implement database connection pooling to reduce connection overhead",
    "Cache frequently accessed book data and catalog metadata",
    "Use read replicas for reporting and analytics queries",
    "Implement request/response compression for large payloads",
    "Add CDN for static assets (book cover images)",
    "Batch notification email sends for efficiency",
    "Use message acknowledgments to guarantee notification delivery"
]

for opt in optimizations:
    doc.add_paragraph(opt, style='List Bullet')

add_heading_style(doc, '8.3 Monitoring and Observability', level=2)
doc.add_paragraph("Production deployment should include:")

monitor_p1 = doc.add_paragraph(style='List Bullet')
monitor_p1.add_run(
    "Distributed tracing (OpenTelemetry or Jaeger) to track requests across services and identify latency bottlenecks."
)

monitor_p2 = doc.add_paragraph(style='List Bullet')
monitor_p2.add_run("Centralized logging (ELK stack or similar) for aggregating logs from all services.")

monitor_p3 = doc.add_paragraph(style='List Bullet')
monitor_p3.add_run("Metrics collection (Prometheus) and dashboards (Grafana) for service health, throughput, and error rates.")

monitor_p4 = doc.add_paragraph(style='List Bullet')
monitor_p4.add_run("Alerting thresholds for SLOs (e.g., payment processing latency < 5s, notification delivery success > 99%).")

doc.add_page_break()

# ===== 9. DESIGN PATTERNS =====
add_heading_style(doc, '9. Design Patterns and Architectural Principles', level=1)

add_heading_style(doc, '9.1 Applied Design Patterns', level=2)

patterns = [
    "Database per Service: Each microservice owns independent database(s). Data sharing occurs through APIs and business identifiers (order_id, customer_id) rather than direct database joins.",
    "API Gateway Pattern: Django monolith acts as primary entry point, aggregating calls to downstream services.",
    "Event-Driven Architecture: Asynchronous events decouple critical services from ancillary workflows (notifications).",
    "Circuit Breaker (implicit): Service calls implement timeouts and graceful fallback on failure.",
    "Domain Events: Services publish domain-specific events (order.paid, shipment.created) enabling loose coupling.",
    "Saga Pattern (manual): Order processing coordinates multiple services through explicit orchestration.",
]

for pattern in patterns:
    doc.add_paragraph(pattern, style='List Bullet')

add_heading_style(doc, '9.2 SOLID Principles', level=2)

solid = [
    "Single Responsibility: Each service handles one business capability (payments, shipping, notifications). Models are focused on domain objects.",
    "Open/Closed: Services are open for extension (new event types, shipping methods) but closed for modification of core logic.",
    "Liskov Substitution: Payment processors can be swapped (Stripe → alternative providers) without changing service interface.",
    "Interface Segregation: Services expose focused APIs (GET /payments/{id}, POST /shipments); clients use only required operations.",
    "Dependency Inversion: Services depend on message broker abstraction for events, not concrete implementations.",
]

for principle in solid:
    doc.add_paragraph(principle, style='List Bullet')

add_heading_style(doc, '9.3 Domain-Driven Design Elements', level=2)
doc.add_paragraph("The system respects DDD principles:")

ddd = [
    "Bounded Contexts: Payment Service, Shipping Service, and Notification Service each represent distinct business domains.",
    "Ubiquitous Language: Core terms (Order, Payment, Shipment) are used consistently across services and APIs.",
    "Aggregates: Order is the aggregate root for payment and shipment operations; Payment and Shipment are separate aggregates.",
    "Domain Events: Services publish events representing significant business occurrences (payment success, shipment creation).",
    "Repository Pattern: Each service manages its persistent entities through database abstraction layer."
]

for item in ddd:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===== 10. DEPLOYMENT WALKTHROUGH =====
add_heading_style(doc, '10. Deployment and Operations Walkthrough', level=1)

add_heading_style(doc, '10.1 Local Development Setup', level=2)

dev_steps = [
    "1. Clone repository and navigate to project root",
    "2. Ensure Docker Desktop is running",
    "3. Execute: docker-compose up",
    "4. Wait 10-15 seconds for database migrations to complete",
    "5. Verify all services are healthy: docker-compose ps",
    "6. Access Django at http://localhost:8000",
    "7. RabbitMQ management UI available at http://localhost:15672 (guest/guest)"
]

for step in dev_steps:
    doc.add_paragraph(step, style='List Number')

add_heading_style(doc, '10.2 Testing the Complete Flow', level=2)
doc.add_paragraph("To verify end-to-end functionality:")

test_steps = [
    "1. Create a customer account via Django web UI (/customers/)",
    "2. Browse books and add items to shopping cart",
    "3. Submit order with payment information",
    "4. Observe payment processing and confirmation of payment.",
    "5. Verify shipment creation with tracking number",
    "6. Check notification service logs for email events published",
    "7. Observe final order status as COMPLETED in Django"
]

for step in test_steps:
    doc.add_paragraph(step, style='List Number')

add_heading_style(doc, '10.3 Monitoring Event Processing', level=2)
doc.add_paragraph("To observe asynchronous events:")

cmd2_p = doc.add_paragraph()
cmd2_run = cmd2_p.add_run("$ docker-compose logs -f notification-service")
cmd2_run.font.name = 'Courier New'
cmd2_run.font.bold = True

doc.add_paragraph(
    "This displays real-time logs of the Notification Service consuming events from the message broker. "
    "Each order.paid and shipment event will be logged as it is processed."
)

doc.add_page_break()

# ===== 11. TESTING =====
add_heading_style(doc, '11. Testing Strategy', level=1)

add_heading_style(doc, '11.1 Current Test Coverage', level=2)
doc.add_paragraph("The system includes functional tests validating critical workflows:")

doc.add_paragraph(
    "Integration tests verify end-to-end order processing including payment and shipping service calls. "
    "Tests exercise the complete flow from order creation through notification publishing. "
    "Services are tested both independently and in integration."
)

add_heading_style(doc, '11.2 Recommended Test Matrix for Production', level=2)

test_matrix = [
    "Unit Tests: Model validation, business logic for price calculation, shipment cost computation",
    "API Integration Tests: Test each service endpoint with valid/invalid inputs, verify error handling",
    "Contract Tests: Verify service-to-service interfaces match expectations (payment API, shipping API)",
    "End-to-End Tests: Full order workflow from Django UI through notification delivery",
    "Resilience Tests: Verify graceful degradation when services are unavailable",
    "Performance Tests: Load test critical paths (order checkout, payment processing)",
    "Security Tests: Test authentication, authorization, input validation, SQL injection prevention"
]

for test in test_matrix:
    doc.add_paragraph(test, style='List Bullet')

add_heading_style(doc, '11.3 Test Execution', level=2)
doc.add_paragraph("Tests can be executed via pytest after system is running:")

cmd3_p = doc.add_paragraph()
cmd3_run = cmd3_p.add_run("$ docker-compose exec web pytest tests/ -v")
cmd3_run.font.name = 'Courier New'
cmd3_run.font.bold = True

doc.add_page_break()

# ===== 12. LIMITATIONS =====
add_heading_style(doc, '12. Current Limitations and Future Roadmap', level=1)

add_heading_style(doc, '12.1 Known Limitations', level=2)

limitations = [
    "Synchronous Payment/Shipping: Order processing blocks on external service calls. Future enhancement: convert to event-driven with eventual consistency and customer polling.",
    "Limited Resilience Patterns: Current timeouts are basic. Production requires circuit breakers, exponential backoff retries, and bulkhead isolation.",
    "No Distributed Transactions: Payment and shipment operations are not strictly ACID across services. Requires saga pattern implementation for failure recovery.",
    "Monolith Remains SPOF: Django application is the gateway for all customer operations. Future: implement API Gateway service for better isolation.",
    "No Service Mesh: Inter-service communication lacks sophisticated routing, load balancing, and mutual TLS. Consider Istio for production.",
    "Manual Authentication: Service-to-service auth uses static bearer tokens. Future: JWT with expiration, automatic rotation, and PKI infrastructure."
]

for limit in limitations:
    doc.add_paragraph(limit, style='List Bullet')

add_heading_style(doc, '12.2 Evolution Roadmap', level=2)

phase1 = doc.add_paragraph(style='List Bullet')
phase1_run = phase1.add_run("Phase 1 (Current): ")
phase1_run.bold = True
phase1.add_run(
    "Functional microservices with Docker Compose, synchronous REST, basic event publishing."
)

phase2 = doc.add_paragraph(style='List Bullet')
phase2_run = phase2.add_run("Phase 2 (Near-term): ")
phase2_run.bold = True
phase2.add_run(
    "Add sophisticated resilience (circuit breaker, retries, bulkhead). "
    "Implement distributed tracing. Add comprehensive API documentation (OpenAPI/Swagger)."
)

phase3 = doc.add_paragraph(style='List Bullet')
phase3_run = phase3.add_run("Phase 3 (Mid-term): ")
phase3_run.bold = True
phase3.add_run(
    "Event-driven order processing with sagas for failure recovery. "
    "Deploy API Gateway separate from web monolith. Add service mesh (Istio) for production networking."
)

phase4 = doc.add_paragraph(style='List Bullet')
phase4_run = phase4.add_run("Phase 4 (Long-term): ")
phase4_run.bold = True
phase4.add_run(
    "Implement CQRS pattern for read/write separation. Add polyglot storage "
    "(MongoDB for notifications, Redis for caching). Migrate web tier to separate frontend (React/Vue). "
    "Implement observability suite (Prometheus, Grafana, Jaeger)."
)

doc.add_page_break()

# ===== 13. DISCUSSION =====
add_heading_style(doc, '13. Discussion', level=1)

add_heading_style(doc, '13.1 Benefits Achieved', level=2)
doc.add_paragraph(
    "The current bookstore implementation demonstrates a pragmatic hybrid architecture: the Django application "
    "retains customer-facing orchestration while specialized FastAPI services handle payment, shipping, and "
    "notification concerns. This separation improves maintainability, keeps operational responsibilities explicit, "
    "and allows the team to evolve service boundaries incrementally rather than through a risky full rewrite."
)

add_heading_style(doc, '13.2 Trade-offs', level=2)
doc.add_paragraph(
    "The architecture introduces cross-service network calls, duplicated operational concerns, and higher deployment "
    "complexity than a pure monolith. Those costs are acceptable for this system because payment and shipping already "
    "have distinct reliability and scaling needs, but they also require stronger observability, retry policies, and "
    "contract discipline than the current academic baseline provides."
)

add_heading_style(doc, '13.3 Academic vs Production Gap', level=2)
doc.add_paragraph(
    "The present solution is suitable as a reference implementation and coursework artifact, but a production rollout "
    "would still require stronger authentication, secrets management, idempotent workflow handling, distributed tracing, "
    "and automated contract testing. The current codebase establishes the right service boundaries; the remaining work is "
    "primarily operational hardening rather than another major architectural redesign."
)

doc.add_page_break()

# ===== 14. CONCLUSION =====
add_heading_style(doc, '14. Conclusion', level=1)

conclusion_text = (
    "The Bookstore microservices architecture successfully balances pragmatism with architectural best practices. "
    "By maintaining a monolithic web tier for customer interaction and order orchestration while delegating specialized "
    "workloads to independent FastAPI services, the system achieves clear separation of concerns and independent scalability.\n\n"
    
    "The event-driven notification workflow demonstrates effective decoupling, enabling the Notification Service to evolve "
    "independently without impacting customer-facing operations. This hybrid approach avoids the operational complexity of "
    "a pure microservices architecture while capturing its benefits: independent deployment, targeted scaling, and team autonomy.\n\n"
    
    "The implementation serves as both a functional e-commerce platform and a reference architecture for teams building "
    "modern Python-based distributed systems. With the roadmap enhancements outlined in Section 12, the system can scale "
    "to production workloads while maintaining the architectural clarity and maintainability that characterizes well-designed "
    "distributed systems.\n\n"
    
    "The Docker Compose-based deployment model enables developers to understand and experiment with the system locally, "
    "while the explicit service boundaries and REST/event-based contracts provide clear paths to cloud-native deployment "
    "with orchestrators like Kubernetes."
)

doc.add_paragraph(conclusion_text)

doc.add_page_break()

# ===== 15. DETAILED SCENARIO WALKTHROUGHS =====
add_heading_style(doc, '15. Detailed Scenario Walkthroughs', level=1)

add_heading_style(doc, '15.1 Order Submission and Payment', level=2)
doc.add_paragraph(
    "A customer completes checkout in the Django storefront, which creates the local order record and then invokes the "
    "Payment Service over HTTP for transactional authorization. On success, the order can transition to a paid state and "
    "continue to downstream fulfillment; on failure, the order remains visible for retry or cancellation handling."
)

add_heading_style(doc, '15.2 Shipment Creation and Tracking', level=2)
doc.add_paragraph(
    "After payment succeeds, the web tier calls the Shipping Service to create a shipment using the order identifier and "
    "delivery details. Shipment status is owned by the shipping database, preserving service autonomy while still allowing "
    "the storefront to display tracking progress through a service client abstraction."
)

add_heading_style(doc, '15.3 Event-Driven Notification Flow', level=2)
doc.add_paragraph(
    "Payment and shipping milestones are published as domain events through RabbitMQ. The Notification Service consumes "
    "those events asynchronously and records outbound notification activity, which decouples user messaging from the "
    "critical checkout path and prevents email delivery latency from blocking transactional operations."
)

add_heading_style(doc, '15.4 Failure Handling Scenario', level=2)
doc.add_paragraph(
    "If payment succeeds but shipment creation fails, the architecture preserves enough separation to retry fulfillment "
    "without re-running payment authorization. This scenario highlights why explicit state transitions, timeout handling, "
    "and eventually compensating actions are required as the system moves closer to production readiness."
)

doc.add_page_break()

# ===== 16. ADR SUMMARY =====
add_heading_style(doc, '16. Architecture Decision Records (ADR Summary)', level=1)

adr_points = [
    "ADR-01: Retain Django as the orchestration layer. The storefront already contains customer-facing workflows, so keeping it as the integration point avoids unnecessary frontend disruption while services are extracted.",
    "ADR-02: Use database-per-service isolation. Payment, shipping, and notification data remain independently owned to prevent schema coupling and to enable separate lifecycle management.",
    "ADR-03: Combine synchronous REST with asynchronous events. Immediate operations such as payment and shipment creation use HTTP, while non-critical notifications use RabbitMQ for loose coupling.",
    "ADR-04: Containerize all components for local reproducibility. Docker Compose provides a consistent development topology and makes the distributed architecture testable on a single workstation.",
]
for point in adr_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_page_break()

# ===== 17. PERFORMANCE AND SCALABILITY =====
add_heading_style(doc, '17. Performance and Scalability Considerations', level=1)

add_heading_style(doc, '17.1 Read-heavy vs Write-heavy Paths', level=2)
doc.add_paragraph(
    "Catalog browsing and book discovery remain comparatively read-heavy, while payment processing and shipment creation "
    "are write-heavy and latency-sensitive. Separating those concerns allows future scaling policies to target the real "
    "bottlenecks rather than scaling the entire application uniformly."
)

add_heading_style(doc, '17.2 Latency Hotspots', level=2)
doc.add_paragraph(
    "The main latency hotspots are synchronous service-to-service calls from the Django application to payment and shipping, "
    "plus downstream event propagation for notifications. Timeouts, connection pooling, and structured retry policies are "
    "therefore more important than raw CPU scaling for the current deployment profile."
)

add_heading_style(doc, '17.3 Data Growth', level=2)
doc.add_paragraph(
    "Order, payment, shipment, and notification records will all grow over time at different rates. Independent databases "
    "make archival, indexing, and retention policies easier to tune per domain, especially once notification history and "
    "operational audit logs become materially larger than core catalog data."
)

doc.add_page_break()

# ===== 18. ROADMAP =====
add_heading_style(doc, '18. Roadmap (Iteration Plan)', level=1)

roadmap_points = [
    "Iteration 1 - Security hardening: Add stronger service authentication, centralize secret management, and enforce authorization checks at integration boundaries.",
    "Iteration 2 - Reliability improvements: Introduce retries, idempotency safeguards, circuit breakers, and better health/readiness reporting across all services.",
    "Iteration 3 - Workflow maturity: Expand event-driven orchestration for order lifecycle transitions and formalize compensating actions for partial failure scenarios.",
    "Iteration 4 - Operational productization: Add tracing, metrics dashboards, CI/CD validation, and production-grade deployment patterns beyond local Docker Compose.",
]
for point in roadmap_points:
    doc.add_paragraph(point, style='List Bullet')

add_heading_style(doc, '18.1 Roadmap Architecture Diagrams', level=2)

diagram_paths = get_roadmap_diagram_paths()
if diagram_paths:
    doc.add_paragraph(
        'The roadmap is illustrated through multiple architecture diagrams, aligned with the staged evolution plan.'
    )
    for index, diagram_path in enumerate(diagram_paths, start=1):
        caption = doc.add_paragraph(f'Figure 18.{index}: {diagram_path.stem}')
        caption.runs[0].italic = True
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(diagram_path), width=Inches(6.2))
        doc.add_paragraph()
else:
    doc.add_paragraph(
        'No rendered roadmap diagrams were found in diagrams/rendered. '
        'Render your PUML files to images named roadmap-*.png (or .jpg/.jpeg) to include them automatically.'
    )

doc.add_page_break()

# ===== APPENDIX A =====
add_heading_style(doc, 'Appendix A - Service List', level=1)

appendix_service_docs = [
    {
        'name': 'api-gateway',
        'get_endpoint': '/books/',
        'get_desc': 'List books through gateway UI.',
        'post_endpoint': '/books/',
        'post_desc': 'Create a book through gateway.',
        'request': '{\n "title": "Clean Architecture",\n "author": "Robert C. Martin",\n "price": "19.99",\n "stock": 20\n}',
    },
    {
        'name': 'customer-service',
        'get_endpoint': '/customers/',
        'get_desc': 'Get all customers.',
        'post_endpoint': '/customers/',
        'post_desc': 'Create customer.',
        'request': '{\n "name": "Alice",\n "email": "alice@example.com"\n}',
    },
    {
        'name': 'staff-service',
        'get_endpoint': '/staffs/',
        'get_desc': 'List staff records.',
        'post_endpoint': '/staffs/',
        'post_desc': 'Create staff record.',
        'request': '{\n "name": "Staff One",\n "email": "staff1@example.com",\n "role": "inventory"\n}',
    },
    {
        'name': 'manager-service',
        'get_endpoint': '/managers/',
        'get_desc': 'List managers.',
        'post_endpoint': '/managers/',
        'post_desc': 'Create manager.',
        'request': '{\n "name": "Manager One",\n "email": "manager1@example.com"\n}',
    },
    {
        'name': 'catalog-service',
        'get_endpoint': '/catalogs/',
        'get_desc': 'List catalogs.',
        'post_endpoint': '/catalogs/',
        'post_desc': 'Create catalog.',
        'request': '{\n "name": "Software Engineering",\n "description": "Books for software practitioners"\n}',
    },
    {
        'name': 'book-service',
        'get_endpoint': '/books/',
        'get_desc': 'List books.',
        'post_endpoint': '/books/',
        'post_desc': 'Create a book.',
        'request': '{\n "title": "Domain-Driven Design",\n "author": "Eric Evans",\n "price": "35.00",\n "stock": 12\n}',
    },
    {
        'name': 'cart-service',
        'get_endpoint': '/carts/',
        'get_desc': 'List carts.',
        'post_endpoint': '/cart-items/',
        'post_desc': 'Add item into cart.',
        'request': '{\n "cart": 1,\n "book_id": 2,\n "quantity": 1\n}',
    },
    {
        'name': 'order-service',
        'get_endpoint': '/orders/',
        'get_desc': 'List orders.',
        'post_endpoint': '/orders/',
        'post_desc': 'Create order.',
        'request': '{\n "customer_id": 1,\n "status": "pending"\n}',
    },
    {
        'name': 'pay-service',
        'get_endpoint': '/payments/',
        'get_desc': 'List payments.',
        'post_endpoint': '/payments/',
        'post_desc': 'Create payment.',
        'request': '{\n "order_id": 1,\n "amount": "59.99",\n "status": "paid"\n}',
    },
    {
        'name': 'ship-service',
        'get_endpoint': '/shipments/',
        'get_desc': 'List shipments.',
        'post_endpoint': '/shipments/',
        'post_desc': 'Create shipment.',
        'request': '{\n "order_id": 1,\n "address": "123 Main Street",\n "status": "created"\n}',
    },
    {
        'name': 'comment-rate-service',
        'get_endpoint': '/comment-rates/',
        'get_desc': 'List comments and ratings.',
        'post_endpoint': '/comment-rates/',
        'post_desc': 'Create a comment/rating.',
        'request': '{\n "customer_id": 1,\n "book_id": 2,\n "comment": "Great read",\n "rating": 5\n}',
    },
    {
        'name': 'recommender-ai-service',
        'get_endpoint': '/recommendations/',
        'get_desc': 'List recommendations.',
        'post_endpoint': '/recommendations/',
        'post_desc': 'Save recommendation score.',
        'request': '{\n "customer_id": 1,\n "book_id": 2,\n "score": 0.93\n}',
    },
]

for index, service in enumerate(appendix_service_docs, start=1):
    add_heading_style(doc, f"{index}. {service['name']}", level=2)
    doc.add_paragraph(f"Base path: `http://{service['name']}:8000`")

    p_get = doc.add_paragraph()
    run_get = p_get.add_run(f"GET {service['get_endpoint']} ")
    run_get.bold = True
    p_get.add_run(f"- Description: {service['get_desc']}")

    p_post = doc.add_paragraph()
    run_post = p_post.add_run(f"POST {service['post_endpoint']} ")
    run_post.bold = True
    p_post.add_run(f"- Description: {service['post_desc']}")

    doc.add_paragraph("Request:")
    request_block = doc.add_paragraph(service['request'])
    for run in request_block.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run(f"Document generated: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}")
footer_run.font.size = Pt(9)
footer_run.font.italic = True
footer_run.font.color.rgb = RGBColor(128, 128, 128)

# Save document
doc_path = Path(r'l:\bookstore\TECHNICAL_REPORT_BOOKSTORE_MICROSERVICES.docx')

try:
    doc.save(doc_path)
    saved_path = doc_path
except PermissionError:
    fallback_name = f"{doc_path.stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{doc_path.suffix}"
    saved_path = doc_path.with_name(fallback_name)
    doc.save(saved_path)

print(f"✓ Technical report (Word document) generated successfully!")
print(f"✓ Location: {saved_path}")
print(f"✓ Format: .docx (fully editable in Microsoft Word)")
print(f"✓ Total sections: 18 major sections aligned to the example report structure")
print(f"✓ Ready for editing and customization!")
