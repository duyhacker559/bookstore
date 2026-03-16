"""
Generate Assignment 05 - Academic Microservice Implementation DOCX report.
Author: Pham Manh Thang - B22DCVT527
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.0)

# ─── Helper: shade table cell ─────────────────────────────────────────────────
def shade_cell(cell, fill_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

# ─── Colour palette ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x37, 0x6C)
TEAL   = RGBColor(0x00, 0x7B, 0x83)
DARK   = RGBColor(0x21, 0x21, 0x21)
GRAY   = RGBColor(0x55, 0x55, 0x55)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG = "F0F4F8"
HEAD_BG = "1A376C"
ALT_BG  = "EAF2FB"

# ─── Base style helpers ───────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, italic=False, color=DARK, name="Calibri"):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12 if level == 2 else 8)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=NAVY, name="Calibri Light")
        # bottom border
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "8")
        bot.set(qn("w:space"), "4")
        bot.set(qn("w:color"), "1A376C")
        pBdr.append(bot)
        pPr.append(pBdr)
    elif level == 2:
        set_font(run, size=13, bold=True, color=TEAL, name="Calibri Light")
    else:
        set_font(run, size=11, bold=True, color=NAVY)
    return p

def add_body(text, italic=False, color=DARK, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, italic=italic, color=color)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_numbered(text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_code(text):
    """Light code-block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)
    # set shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  CODE_BG)
    pPr.append(shd)
    return p

def add_info_box(label: str, body: str):
    """Coloured two-cell table used as an info-box."""
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, ALT_BG)
    cell.paragraphs[0].clear()
    lrun = cell.paragraphs[0].add_run(label + "  ")
    set_font(lrun, bold=True, color=NAVY, size=10)
    brun = cell.paragraphs[0].add_run(body)
    set_font(brun, size=10, color=DARK)
    doc.add_paragraph()

def styled_api_table(rows_data):
    """rows_data: list of (method, endpoint, description)"""
    tbl = doc.add_table(rows=1 + len(rows_data), cols=3)
    tbl.style = "Table Grid"
    # Header
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(["Method / Endpoint", "Path", "Description"]):
        shade_cell(hdr[i], HEAD_BG)
        p  = hdr[i].paragraphs[0]
        p.clear()
        r  = p.add_run(txt)
        set_font(r, bold=True, color=WHITE, size=10)
    for ri, (method, path, desc) in enumerate(rows_data, start=1):
        cells = tbl.rows[ri].cells
        if ri % 2 == 0:
            shade_cell(cells[0], ALT_BG)
            shade_cell(cells[1], ALT_BG)
            shade_cell(cells[2], ALT_BG)
        for ci, txt in enumerate([method, path, desc]):
            p = cells[ci].paragraphs[0]
            p.clear()
            r = p.add_run(txt)
            if ci == 0:
                set_font(r, bold=True, size=9, color=TEAL, name="Courier New")
            elif ci == 1:
                set_font(r, size=9, color=NAVY, name="Courier New")
            else:
                set_font(r, size=9.5)
    # column widths
    for ri in range(len(tbl.rows)):
        tbl.rows[ri].cells[0].width = Cm(3.0)
        tbl.rows[ri].cells[1].width = Cm(6.0)
        tbl.rows[ri].cells[2].width = Cm(8.0)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
cover_tbl = doc.add_table(rows=1, cols=1)
cover_tbl.style = "Table Grid"
cc = cover_tbl.rows[0].cells[0]
shade_cell(cc, "1A376C")
cc.width = Inches(6.5)

def cover_line(cell, text, sz, bold=False, space_before=0, space_after=6, color=WHITE):
    p   = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=sz, bold=bold, color=color, name="Calibri Light")
    return p

cc.paragraphs[0].clear()

cover_line(cc, "POSTS & TELECOMMUNICATIONS", 11, space_before=20)
cover_line(cc, "INSTITUTE OF TECHNOLOGY", 11)
cover_line(cc, "Faculty of Information Technology", 10,
           color=RGBColor(0xAD, 0xD8, 0xE6))
cover_line(cc, "─" * 52, 9, color=RGBColor(0x7B, 0xB2, 0xD9), space_after=24)

cover_line(cc, "ASSIGNMENT 05", 22, bold=True, space_before=10)
cover_line(cc, "Academic Microservice Implementation", 16,
           color=RGBColor(0xAD, 0xD8, 0xE6), space_after=4)
cover_line(cc, "BookStore System — Full Technical Report", 13,
           color=RGBColor(0xFF, 0xFF, 0xFF), space_after=30)

cover_line(cc, "─" * 52, 9, color=RGBColor(0x7B, 0xB2, 0xD9), space_after=18)
cover_line(cc, "Subject:  Software Architecture & Microservices", 11)
cover_line(cc, "Student:  Pham Manh Thang", 11)
cover_line(cc, "Student ID:  B22DCVT527", 11)
cover_line(cc, "Date Submitted:  April 4, 2025", 11, space_after=30)
cover_line(cc, "─" * 52, 9, color=RGBColor(0x7B, 0xB2, 0xD9), space_after=30)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS  (manual)
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Table of Contents", 1)

toc_items = [
    ("1", "Introduction",                                    ""),
    ("  1.1", "Background",                                  ""),
    ("  1.2", "Problem Statement",                           ""),
    ("  1.3", "Objective",                                   ""),
    ("2", "Assignment Requirements",                         ""),
    ("  2.1", "Objectives (§4.1)",                           ""),
    ("  2.2", "Required Services (§4.2)",                    ""),
    ("  2.3", "Functional Requirements (§4.3)",              ""),
    ("  2.4", "Technical Requirements (§4.4)",               ""),
    ("  2.5", "Deliverables (§4.5)",                         ""),
    ("3", "From Monolith to Microservices",                  ""),
    ("4", "Design Principles",                               ""),
    ("5", "DDD Decomposition",                               ""),
    ("6", "Target Architecture Overview",                    ""),
    ("7", "Service Implementation Summary",                  ""),
    ("8", "API Documentation",                               ""),
    ("9", "Functional Requirement Coverage",                 ""),
    ("10", "Deployment, Build & Execution",                  ""),
    ("11", "Reliability & Resilience Assessment",            ""),
    ("12", "Security Considerations",                        ""),
    ("13", "Testing & Verification",                         ""),
    ("14", "Scenario Walkthroughs",                          ""),
    ("15", "Architecture Decision Records",                  ""),
    ("16", "Performance & Scalability",                      ""),
    ("17", "Roadmap (Iteration Plan)",                       ""),
    ("18", "Discussion & Conclusion",                        ""),
    ("Appendix A", "Service List",                           ""),
    ("Appendix B", "Common HTTP Status Codes",               ""),
]
for num, title, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    indent = 0.4 if num.startswith("  ") else 0.0
    p.paragraph_format.left_indent = Cm(indent)
    tab_stops = p.paragraph_format.tab_stops
    run = p.add_run(f"{num.strip()}    {title}")
    set_font(run, size=10 if num.startswith("  ") else 10.5,
             bold=not num.startswith("  "), color=NAVY if not num.startswith("  ") else DARK)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("1. Introduction", 1)

add_heading("1.1 Background", 2)
add_body(
    "Traditional monolithic systems package all modules into one deployable artifact and "
    "typically share a single database. While this model simplifies initial development, growth "
    "introduces technical and organisational friction. In a BookStore context, identity, "
    "catalogue, cart, order, payment, shipping, review, and recommendation evolve at different "
    "speeds and have distinct runtime profiles."
)

add_heading("1.2 Problem Statement", 2)
add_body("A monolith for BookStore causes:")
add_bullet("Hard scaling for specific hot paths (e.g., catalogue browsing).")
add_bullet("Risky deployment — all modules are released together.")
add_bullet("Tight coupling across teams and code modules.")
add_bullet("Technology lock-in and lower parallel development throughput.")

add_heading("1.3 Objective", 2)
add_body("The objective is to implement an academic yet realistic microservice architecture that:")
add_bullet("Decomposes domain capabilities into independent services.")
add_bullet("Enforces database-per-service boundaries.")
add_bullet("Exposes API-first REST contracts.")
add_bullet("Satisfies all assignment functional requirements.")
add_bullet("Runs on Docker Compose and supports service-level testing.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. ASSIGNMENT REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("2. Assignment Requirements", 1)

add_heading("2.1 Objective (§4.1)", 2)
add_body(
    "Students must decompose a monolithic BookStore into microservices using "
    "Django REST Framework."
)

add_heading("2.2 Required Services (§4.2)", 2)
services = [
    "staff-service", "manager-service", "customer-service", "catalog-service",
    "book-service", "cart-service", "order-service", "ship-service",
    "pay-service", "comment-rate-service", "recommender-ai-service", "api-gateway",
]
for i, svc in enumerate(services, 1):
    add_numbered(svc)

add_heading("2.3 Functional Requirements (§4.3)", 2)
frs = [
    "Customer registration automatically creates a cart.",
    "Staff manages books.",
    "Customer adds books to cart, views cart, updates cart.",
    "Order triggers payment and shipping; customer selects pay and ship method.",
    "Customer can rate books.",
]
for fr in frs:
    add_numbered(fr)

add_heading("2.4 Technical Requirements (§4.4)", 2)
tech_reqs = [
    "Django REST Framework",
    "REST inter-service calls",
    "Docker Compose",
    "Independent databases per service",
]
for tr in tech_reqs:
    add_bullet(tr)

add_heading("2.5 Deliverables (§4.5)", 2)
deliverables = [
    "GitHub repository with all service source code.",
    "Architecture diagram for each service.",
    "API documentation.",
    "10-minute demo video.",
    "8–12 page technical report.",
]
for d in deliverables:
    add_numbered(d)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. FROM MONOLITH TO MICROSERVICES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("3. From Monolith to Microservices", 1)

add_heading("3.1 Monolithic Baseline", 2)
add_body("Monolith characteristics relevant to BookStore:")
add_bullet("Single codebase, single runtime process.")
add_bullet("Shared database schema across all modules.")
add_bullet("In-process module calls (no network boundary).")
add_bullet("One deployment pipeline for all features.")
add_body("Advantages at small scale:", italic=True)
add_bullet("Simple local setup and debugging.")
add_bullet("Easier ACID transactions spanning multiple modules.")

add_heading("3.2 Monolith Limitations", 2)
add_body("As scope increases, practical issues appear:")
add_bullet("Scaling mismatch: checkout and catalogue traffic differ; monolith scales all-or-nothing.")
add_bullet("Release risk: a change in the review module can impact the checkout deployment.")
add_bullet("Team contention: merge conflicts and long integration cycles.")
add_bullet("Schema coupling: changing one bounded context may break another query path.")

add_heading("3.3 Microservice Transition Rationale", 2)
add_body(
    "Microservices are selected to isolate business capabilities and operational concerns. "
    "The BookStore domain naturally fits service decomposition because each context "
    "(identity, catalogue, ordering, payment, shipping, review, recommendation) can be "
    "modelled with explicit API contracts and independent persistence."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. DESIGN PRINCIPLES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("4. Design Principles Applied", 1)

principles = [
    (
        "4.1 Single Responsibility at Service Level",
        "Each service owns one business capability:\n"
        "• customer-service handles customer records.\n"
        "• cart-service handles carts and cart items.\n"
        "• pay-service handles payment records.\n"
        "This reduces cognitive load and allows focused testing."
    ),
    (
        "4.2 Database per Service",
        "No service directly reads another service's database. Data sharing is performed by "
        "API calls and business identifiers (customer_id, book_id, order_id). This supports "
        "autonomy and independent schema evolution."
    ),
    (
        "4.3 API-First Design",
        "The system exposes explicit REST endpoints for each service. The current version "
        "focuses on GET and POST contracts required by assignment scenarios, with a pathway "
        "to full CRUD and OpenAPI publication."
    ),
    (
        "4.4 Loose Coupling",
        "Dependencies are explicit and narrow: customer-service → cart-service (cart auto-"
        "creation), cart-service → book-service (book validation), gateway → downstream "
        "services for UI rendering."
    ),
    (
        "4.5 Stateless Services",
        "Service endpoints are stateless in request processing (no in-memory session "
        "dependency). Authentication with JWT is listed as a principle; the current academic "
        "implementation can add JWT in the next iteration."
    ),
    (
        "4.6 Resilience",
        "Current code includes basic request timeout for selected calls (timeout=5 s). Full "
        "resilience stack (retry, circuit breaker, fallback policy) is proposed in the roadmap."
    ),
]
for title, body in principles:
    add_heading(title, 2)
    for line in body.split("\n"):
        if line.startswith("•"):
            add_bullet(line[1:].strip())
        else:
            add_body(line)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. DDD DECOMPOSITION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("5. DDD Decomposition", 1)

add_heading("5.1 Bounded Context Mapping", 2)
contexts = [
    ("Identity context",        "customer-service, staff-service, manager-service"),
    ("Catalogue context",       "catalog-service, book-service"),
    ("Ordering context",        "cart-service, order-service"),
    ("Payment context",         "pay-service"),
    ("Shipping context",        "ship-service"),
    ("Review context",          "comment-rate-service"),
    ("Recommendation context",  "recommender-ai-service"),
    ("Access context",          "api-gateway"),
]
tbl = doc.add_table(rows=1 + len(contexts), cols=2)
tbl.style = "Table Grid"
for ci, h in enumerate(["Bounded Context", "Services"]):
    shade_cell(tbl.rows[0].cells[ci], HEAD_BG)
    p = tbl.rows[0].cells[ci].paragraphs[0]
    p.clear()
    r = p.add_run(h)
    set_font(r, bold=True, color=WHITE, size=10)
for ri, (ctx, svcs) in enumerate(contexts, 1):
    cells = tbl.rows[ri].cells
    if ri % 2 == 0:
        shade_cell(cells[0], ALT_BG)
        shade_cell(cells[1], ALT_BG)
    for ci, txt in enumerate([ctx, svcs]):
        p = cells[ci].paragraphs[0]
        p.clear()
        r = p.add_run(txt)
        set_font(r, size=10, bold=(ci == 0))
doc.add_paragraph()

add_heading("5.2 Ubiquitous Language", 2)
add_body("Core domain terms used consistently in APIs and models:")
for term in ["Customer, Staff, Manager", "Catalog, Book",
             "Cart, CartItem, Order",
             "Payment, Shipment", "CommentRate, Recommendation"]:
    add_bullet(term)

add_heading("5.3 Context Mapping (Integration Style)", 2)
add_bullet("Upstream identity → ordering: customer creation triggers cart creation.")
add_bullet("Upstream catalogue → ordering: cart validates book before insertion.")
add_bullet("Ordering → payment/shipping: order references downstream operational records.")

add_heading("5.4 Aggregate and Data Ownership", 2)
add_bullet("Cart aggregate owns CartItem child collection.")
add_bullet("Order aggregate remains isolated from payment/shipment storage.")
add_bullet("Recommendation aggregate stores model outputs independently of transactional tables.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. TARGET ARCHITECTURE OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("6. Target Architecture Overview", 1)

add_heading("6.1 System Topology", 2)
add_bullet("External clients connect to API Gateway (port 8000).")
add_bullet("API Gateway exposes user-facing routes and forwards/aggregates downstream data.")
add_bullet("11 domain services handle business APIs on internal ports (8001–8011).")
add_bullet("Each service has private persistence (SQLite for academic; PostgreSQL for production).")

add_heading("6.2 Communication Pattern", 2)
add_body("Synchronous REST over HTTP (current iteration).")
add_body("Logical cross-service flows:")
add_bullet("cart-service validates book existence via book-service GET /books/.")
add_bullet("customer-service auto-creates cart via cart-service POST /carts/.")
add_bullet("Recommendation and review services use cross-context identifiers (customer_id, book_id).")

add_heading("6.3 Deployment Pattern", 2)
add_bullet("Each service is containerised with its own Dockerfile.")
add_bullet("Orchestrated via docker-compose.yml.")
add_bullet("Enables independent restart and selective scaling per service.")

add_heading("6.4 Architecture Quality Attributes", 2)
qa_rows = [
    ("Scalability",     "Selective scale-out by service role (e.g., catalog-service during browse peaks)"),
    ("Maintainability", "Bounded codebases per domain — easier to navigate and test"),
    ("Reliability",     "Service failure impact is localised; other services continue operating"),
    ("Evolvability",    "API contract boundaries support iterative schema and logic changes"),
]
tbl = doc.add_table(rows=1 + len(qa_rows), cols=2)
tbl.style = "Table Grid"
for ci, h in enumerate(["Quality Attribute", "How Achieved"]):
    shade_cell(tbl.rows[0].cells[ci], HEAD_BG)
    p = tbl.rows[0].cells[ci].paragraphs[0]
    p.clear()
    r = p.add_run(h)
    set_font(r, bold=True, color=WHITE, size=10)
for ri, (qa, how) in enumerate(qa_rows, 1):
    cells = tbl.rows[ri].cells
    if ri % 2 == 0:
        shade_cell(cells[0], ALT_BG); shade_cell(cells[1], ALT_BG)
    for ci, txt in enumerate([qa, how]):
        p = cells[ci].paragraphs[0]; p.clear()
        r = p.add_run(txt)
        set_font(r, size=10, bold=(ci == 0))
doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. SERVICE IMPLEMENTATION SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("7. Service Implementation Summary", 1)

add_heading("7.1 API Gateway", 2)
add_body("Responsibilities:")
add_bullet("Single entry point for UI templates and service navigation.")
add_bullet("Proxy and aggregate calls to downstream APIs.")
add_bullet("Health endpoint for liveness check.")
add_body("Notable routes: /books/, /customers/, /cart-items/, /cart-detail/, /services/, /health/")

add_heading("7.2 Identity Services", 2)
identity_svcs = [
    ("customer-service",
     "Model: Customer(name, email unique)\n"
     "API: GET/POST /customers/\n"
     "Side effect: auto-calls cart-service POST /carts/ after customer creation."),
    ("staff-service",
     "Model: Staff(name, email unique, role)\n"
     "API: GET/POST /staffs/"),
    ("manager-service",
     "Model: Manager(name, email unique)\n"
     "API: GET/POST /managers/"),
]
for svc, detail in identity_svcs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(svc)
    set_font(r, bold=True, color=TEAL, size=11)
    for line in detail.split("\n"):
        add_body(line, indent=True)

add_heading("7.3 Catalogue Services", 2)
cat_svcs = [
    ("catalog-service",
     "Model: Catalog(name, description)\n"
     "API: GET/POST /catalogs/"),
    ("book-service",
     "Model: Book(title, author, price, stock)\n"
     "API: GET/POST /books/"),
]
for svc, detail in cat_svcs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(svc)
    set_font(r, bold=True, color=TEAL, size=11)
    for line in detail.split("\n"):
        add_body(line, indent=True)

add_heading("7.4 Ordering Services", 2)
order_svcs = [
    ("cart-service",
     "Models: Cart(customer_id), CartItem(cart, book_id, quantity)\n"
     "APIs: GET/POST /carts/, GET /carts/{id}/, POST /cart-items/\n"
     "Validation: queries book-service /books/ and rejects unknown book_id."),
    ("order-service",
     "Model: Order(customer_id, status)\n"
     "API: GET/POST /orders/"),
]
for svc, detail in order_svcs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(svc)
    set_font(r, bold=True, color=TEAL, size=11)
    for line in detail.split("\n"):
        add_body(line, indent=True)

add_heading("7.5 Payment and Shipping", 2)
pay_svcs = [
    ("pay-service",
     "Model: Payment(order_id, amount, status)\n"
     "API: GET/POST /payments/"),
    ("ship-service",
     "Model: Shipment(order_id, address, status)\n"
     "API: GET/POST /shipments/"),
]
for svc, detail in pay_svcs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(svc)
    set_font(r, bold=True, color=TEAL, size=11)
    for line in detail.split("\n"):
        add_body(line, indent=True)

add_heading("7.6 Review and Recommendation", 2)
rev_svcs = [
    ("comment-rate-service",
     "Model: CommentRate(customer_id, book_id, comment, rating)\n"
     "API: GET/POST /comment-rates/"),
    ("recommender-ai-service",
     "Model: Recommendation(customer_id, book_id, score)\n"
     "API: GET/POST /recommendations/"),
]
for svc, detail in rev_svcs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(svc)
    set_font(r, bold=True, color=TEAL, size=11)
    for line in detail.split("\n"):
        add_body(line, indent=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. API DOCUMENTATION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("8. API Documentation", 1)
add_body(
    "Base URL: http://localhost:8000 (API Gateway). Internal service ports are mapped "
    "by Docker Compose. All services accept and return JSON. Current version provides list "
    "and create endpoints (GET, POST)."
)

# ── 8.1 API Gateway ─────────────────────────────────────────────────────────
add_heading("8.1 API Gateway  (port 8000)", 2)
styled_api_table([
    ("GET",      "/",                          "Render home page (HTML)"),
    ("GET|POST", "/books/",                    "View / create books via gateway UI"),
    ("GET|POST", "/customers/",                "View / create customers; triggers cart creation"),
    ("GET|POST", "/cart-items/",               "Add item to a cart via gateway"),
    ("GET|POST", "/cart-detail/",              "Enter cart_id and view cart detail"),
    ("GET",      "/services/",                 "Show links to all downstream service pages"),
    ("GET",      "/services/{service_key}/",   "Generic list page for one downstream service"),
    ("GET",      "/health/",                   "Gateway liveness probe — returns 'ok'"),
])

# ── 8.2 customer-service ────────────────────────────────────────────────────
add_heading("8.2 customer-service  (internal: customer-service:8000)", 2)
styled_api_table([
    ("GET",  "/customers/",  "List all customers"),
    ("POST", "/customers/",  "Create customer; side effect: POST /carts/ on cart-service"),
])
add_body("Sample POST payload:")
for line in ['{"name": "Pham Manh Thang", "email": "thang@example.com"}']:
    add_code(line)

# ── 8.3 staff-service ───────────────────────────────────────────────────────
add_heading("8.3 staff-service  (internal: staff-service:8000)", 2)
styled_api_table([
    ("GET",  "/staffs/", "List staff records"),
    ("POST", "/staffs/", "Create staff record (fields: name, email, role)"),
])
add_code('{"name": "Staff A", "email": "staff@example.com", "role": "inventory"}')

# ── 8.4 manager-service ─────────────────────────────────────────────────────
add_heading("8.4 manager-service  (internal: manager-service:8000)", 2)
styled_api_table([
    ("GET",  "/managers/", "List managers"),
    ("POST", "/managers/", "Create manager (fields: name, email)"),
])
add_code('{"name": "Manager A", "email": "manager@example.com"}')

# ── 8.5 catalog-service ─────────────────────────────────────────────────────
add_heading("8.5 catalog-service  (internal: catalog-service:8000)", 2)
styled_api_table([
    ("GET",  "/catalogs/", "List catalogues"),
    ("POST", "/catalogs/", "Create catalogue (fields: name, description)"),
])
add_code('{"name": "Software Engineering", "description": "Books about architecture and design"}')

# ── 8.6 book-service ────────────────────────────────────────────────────────
add_heading("8.6 book-service  (internal: book-service:8000)", 2)
styled_api_table([
    ("GET",  "/books/", "List books"),
    ("POST", "/books/", "Create book (fields: title, author, price, stock)"),
])
add_code('{"title": "Clean Architecture", "author": "Robert C. Martin", "price": "19.99", "stock": 20}')

# ── 8.7 cart-service ────────────────────────────────────────────────────────
add_heading("8.7 cart-service  (internal: cart-service:8000)", 2)
styled_api_table([
    ("GET",  "/carts/",           "List carts"),
    ("POST", "/carts/",           "Create a cart   {customer_id: int}"),
    ("GET",  "/carts/{cart_id}/", "Retrieve one cart with its items"),
    ("POST", "/cart-items/",      "Add item to cart; validates book_id against book-service"),
])
add_code('POST /cart-items/  →  {"cart": 1, "book_id": 2, "quantity": 1}')
add_body("Validation: calls GET http://book-service:8000/books/ — returns 404 if book_id absent.")

# ── 8.8 order-service ───────────────────────────────────────────────────────
add_heading("8.8 order-service  (internal: order-service:8000)", 2)
styled_api_table([
    ("GET",  "/orders/", "List orders"),
    ("POST", "/orders/", "Create order  {customer_id, status}"),
])
add_code('{"customer_id": 1, "status": "pending"}')

# ── 8.9 pay-service ─────────────────────────────────────────────────────────
add_heading("8.9 pay-service  (internal: pay-service:8000)", 2)
styled_api_table([
    ("GET",  "/payments/", "List payments"),
    ("POST", "/payments/", "Create payment  {order_id, amount, status}"),
])
add_code('{"order_id": 1, "amount": "99.50", "status": "pending"}')

# ── 8.10 ship-service ───────────────────────────────────────────────────────
add_heading("8.10 ship-service  (internal: ship-service:8000)", 2)
styled_api_table([
    ("GET",  "/shipments/", "List shipments"),
    ("POST", "/shipments/", "Create shipment  {order_id, address, status}"),
])
add_code('{"order_id": 1, "address": "Ha Noi", "status": "preparing"}')

# ── 8.11 comment-rate-service ───────────────────────────────────────────────
add_heading("8.11 comment-rate-service  (internal: comment-rate-service:8000)", 2)
styled_api_table([
    ("GET",  "/comment-rates/", "List comments and ratings"),
    ("POST", "/comment-rates/", "Create comment/rating  {customer_id, book_id, comment, rating}"),
])
add_code('{"customer_id": 1, "book_id": 2, "comment": "Good book", "rating": 5}')

# ── 8.12 recommender-ai-service ─────────────────────────────────────────────
add_heading("8.12 recommender-ai-service  (internal: recommender-ai-service:8000)", 2)
styled_api_table([
    ("GET",  "/recommendations/", "List recommendations"),
    ("POST", "/recommendations/", "Save recommendation score  {customer_id, book_id, score}"),
])
add_code('{"customer_id": 1, "book_id": 2, "score": 0.93}')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. FUNCTIONAL REQUIREMENT COVERAGE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("9. Functional Requirement Coverage", 1)

frc = [
    (
        "FR1 — Customer Registration Automatically Creates a Cart",
        "Implemented in customer-service POST /customers/.",
        [
            "Save customer record in customer-db.",
            "Internally call cart-service POST /carts/ with {customer_id}.",
            "Return 201 Created to the caller.",
            "Postcondition: customer record in customer-db; cart record in cart-db.",
        ],
        "Gap: cart call is not yet wrapped in compensation logic. "
        "Improvement: transactional outbox + event-driven cart creation."
    ),
    (
        "FR2 — Staff Manages Books",
        "Current implementation provides service-level CRUD foundation.",
        [
            "book-service POST /books/ creates books.",
            "Gateway exposes /books/ page for manual operation.",
        ],
        "Gap: role-based authorisation (staff only) can be added via JWT and policy middleware."
    ),
    (
        "FR3 — Customer Adds Books to Cart, Views Cart, Updates Cart",
        "Covered for add and view flows.",
        [
            "Add item: cart-service POST /cart-items/.",
            "View cart: cart-service GET /carts/{id}/.",
            "Book existence validated via book-service call.",
        ],
        "Gap: explicit update endpoint for cart-item quantity is not yet implemented; proposed in roadmap."
    ),
    (
        "FR4 — Order Triggers Payment and Shipping",
        "APIs exist for order/payment/shipping services.",
        [
            "order-service manages order entities.",
            "pay-service and ship-service manage operational records.",
        ],
        "Gap: automatic orchestration from order creation to payment/shipment is currently a logical "
        "flow in the diagram; can be extended with orchestration or event pattern."
    ),
    (
        "FR5 — Customer Can Rate Books",
        "Implemented via comment-rate-service.",
        [
            "POST /comment-rates/ stores rating and comment with customer/book references.",
        ],
        ""
    ),
]
for title, intro, steps, gap in frc:
    add_heading(title, 2)
    add_body(intro)
    for s in steps:
        add_bullet(s)
    if gap:
        add_body("Note: " + gap, italic=True, color=GRAY)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. DEPLOYMENT, BUILD & EXECUTION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("10. Deployment, Build & Execution", 1)

add_heading("10.1 Containerisation", 2)
add_body("Each service includes:")
add_bullet("Independent Dockerfile.")
add_bullet("requirements.txt scoped to the service.")
add_bullet("Its own Django project and app.")

add_heading("10.2 Orchestration with Docker Compose", 2)
add_body("docker-compose.yml coordinates service startup, networking, and port mapping. Benefits:")
add_bullet("Consistent and reproducible environment.")
add_bullet("Repeatable demo flow.")
add_bullet("Simple per-service restart and log observation.")

add_heading("10.3 Starting the System", 2)
add_code("# Build and start all services")
add_code("docker compose up --build")
add_code("")
add_code("# Start in detached mode")
add_code("docker compose up -d --build")
add_code("")
add_code("# Scale a specific service")
add_code("docker compose up --scale catalog-service=3")

add_heading("10.4 Verifying Services", 2)
add_code("# API Gateway health")
add_code("curl http://localhost:8000/health/")
add_code("")
add_code("# Create a customer (triggers cart creation)")
add_code('curl -X POST http://localhost:8000/customers/ \\')
add_code('  -H "Content-Type: application/json" \\')
add_code('  -d \'{"name": "Pham Manh Thang", "email": "thang@example.com"}\'')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 11. RELIABILITY & RESILIENCE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("11. Reliability & Resilience Assessment", 1)

add_heading("11.1 Current State", 2)
add_bullet("Timeout in selected HTTP requests (timeout=5 s) in gateway and customer/cart service calls.")
add_bullet("Service-level isolation limits blast radius of individual failures.")

add_heading("11.2 Gaps", 2)
add_bullet("No retry policy for transient downstream failures.")
add_bullet("No circuit breaker to prevent cascading failures.")
add_bullet("No message broker for async decoupling in critical workflows.")
add_bullet("No idempotency key handling for repeated client submissions.")

add_heading("11.3 Priority Improvement Plan", 2)
improvements = [
    "Add retries with exponential back-off for safe idempotent GET requests.",
    "Introduce circuit breaker around inter-service calls.",
    "Add consistent request timeout and fallback responses across all services.",
    "Introduce event bus (RabbitMQ/Kafka) for order-created workflows.",
    "Add distributed tracing (OpenTelemetry) and correlation IDs.",
]
for i, imp in enumerate(improvements, 1):
    add_numbered(imp)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 12. SECURITY CONSIDERATIONS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("12. Security Considerations", 1)

add_heading("12.1 Current State", 2)
add_bullet("Basic service exposure for academic environment.")
add_bullet("No unified auth token enforcement yet.")

add_heading("12.2 Recommended Security Controls", 2)
sec_items = [
    "JWT-based auth at gateway with service-side validation.",
    "Role-based authorisation for staff/manager operations (RBAC).",
    "Input validation hardening for all payload fields (serializer constraints).",
    "Rate limiting and abuse protection on gateway routes.",
    "Secrets management for credentials and service configs (e.g., Docker secrets, Vault).",
    "HTTPS/TLS termination at gateway for production deployments.",
]
for s in sec_items:
    add_bullet(s)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 13. TESTING & VERIFICATION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("13. Testing & Verification", 1)

add_heading("13.1 Functional Tests Executed", 2)
add_body("Core flows verified:")
tests_done = [
    "Customer creation and auto-cart side effect.",
    "Book creation and listing via book-service.",
    "Cart item insertion with book validation.",
    "Cart detail retrieval including items.",
    "Gateway route availability and service list rendering.",
]
for t in tests_done:
    add_bullet(t)

add_heading("13.2 Automated Test Matrix (Recommended)", 2)
test_matrix = [
    ("Unit tests",       "Serializer/model validation per service"),
    ("Integration tests","Inter-service call flows (e.g., customer → cart)"),
    ("Contract tests",   "Gateway ↔ service endpoint contracts"),
    ("Negative tests",   "Invalid IDs, malformed payloads, duplicate email"),
    ("Resilience tests", "Timeout and downstream unavailability scenarios"),
]
tbl = doc.add_table(rows=1 + len(test_matrix), cols=2)
tbl.style = "Table Grid"
for ci, h in enumerate(["Test Type", "Coverage Target"]):
    shade_cell(tbl.rows[0].cells[ci], HEAD_BG)
    p = tbl.rows[0].cells[ci].paragraphs[0]; p.clear()
    r = p.add_run(h); set_font(r, bold=True, color=WHITE, size=10)
for ri, (tt, cov) in enumerate(test_matrix, 1):
    cells = tbl.rows[ri].cells
    if ri % 2 == 0:
        shade_cell(cells[0], ALT_BG); shade_cell(cells[1], ALT_BG)
    for ci, txt in enumerate([tt, cov]):
        p = cells[ci].paragraphs[0]; p.clear()
        r = p.add_run(txt); set_font(r, size=10, bold=(ci == 0))
doc.add_paragraph()

add_heading("13.3 Acceptance Criteria Mapping", 2)
ac = [
    ("Diagram per service", "Completed and explained in Section 6"),
    ("API documentation",   "Completed in Section 8"),
    ("Technical report",    "This document"),
    ("GitHub repository",   "Operational task"),
    ("Demo video",          "Operational task (10-minute walkthrough)"),
]
for criterion, status in ac:
    add_bullet(f"{criterion}: {status}")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 14. SCENARIO WALKTHROUGHS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("14. End-to-End Scenario Walkthroughs", 1)

add_heading("14.1 Scenario A — Customer Registration with Automatic Cart Creation", 2)
add_body("Preconditions: customer-service and cart-service are healthy.")
for step in [
    "Client submits POST /customers/ with name and email.",
    "customer-service validates payload and inserts new customer row.",
    "customer-service sends POST /carts/ to cart-service with {customer_id}.",
    "cart-service creates a cart owned by this customer.",
    "customer-service returns 201 to caller.",
]:
    add_numbered(step)
add_body(
    "Postconditions: new customer in customer-db; new cart in cart-db with matching customer_id.",
    italic=True, color=GRAY
)

add_heading("14.2 Scenario B — Add Book to Cart with Book Validation", 2)
add_body("Preconditions: book exists in book-service; cart exists in cart-service.")
for step in [
    "Client submits POST /cart-items/ with cart, book_id, quantity.",
    "cart-service calls GET /books/ on book-service.",
    "cart-service checks whether book_id exists in the response list.",
    "If valid, saves CartItem; otherwise returns 404 Book not found.",
]:
    add_numbered(step)
add_body(
    "Improvement option: GET /books/{id} endpoint and response caching.",
    italic=True, color=GRAY
)

add_heading("14.3 Scenario C — Order, Payment and Shipment Lifecycle", 2)
add_body(
    "Current behaviour: order-service, pay-service, and ship-service are separate APIs "
    "that can be manually sequenced."
)
add_body("Recommended orchestrated flow (next iteration):")
for step in [
    "Create order (status=pending).",
    "Trigger payment creation.",
    "If payment succeeds, create shipment.",
    "Update order status to processing then completed.",
]:
    add_numbered(step)
add_body(
    "Recommended design: async domain events (OrderCreated, PaymentSucceeded, ShipmentCreated) "
    "to reduce runtime coupling.",
    italic=True, color=GRAY
)

add_heading("14.4 Scenario D — Rating and Recommendation", 2)
for step in [
    "Customer posts review to comment-rate-service.",
    "Recommendation service periodically reads interaction features and writes scores.",
    "Gateway exposes recommendation list for UI consumption.",
]:
    add_numbered(step)
add_body(
    "Benefit: separation of transactional ordering and analytical recommendation; "
    "recommendation model evolution does not require checkout service deployment.",
    italic=True, color=GRAY
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 15. ARCHITECTURE DECISION RECORDS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("15. Architecture Decision Records (ADR Summary)", 1)

adrs = [
    (
        "ADR-01: API Gateway as Single Public Entry Point",
        "External clients access only gateway routes.",
        "Simplifies client integration and internal endpoint evolution. Enables future centralised "
        "concerns: auth, throttling, logging.",
        "Gateway becomes a critical component and must be monitored; single point of failure risk."
    ),
    (
        "ADR-02: Database per Service",
        "Every service owns private persistence.",
        "Avoids tight schema coupling and cross-team migration conflicts.",
        "Cross-service joins are impossible at DB level — must be composed via APIs/events."
    ),
    (
        "ADR-03: Synchronous REST for Initial Integration",
        "Use synchronous REST over HTTP for first academic iteration.",
        "Simpler implementation and debugging for coursework timeline.",
        "Higher runtime coupling than event-driven approach; needs resilience controls for production."
    ),
    (
        "ADR-04: Stateless Services",
        "No server-side user session required by service logic.",
        "Better horizontal scaling and simpler failover.",
        "Authentication context should move to JWT token in the next iteration."
    ),
]
for title, decision, reasoning, consequences in adrs:
    add_heading(title, 2)
    add_info_box("Decision:",     decision)
    add_info_box("Reasoning:",    reasoning)
    add_info_box("Consequences:", consequences)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 16. PERFORMANCE & SCALABILITY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("16. Performance & Scalability Considerations", 1)

add_heading("16.1 Read-heavy vs Write-heavy Paths", 2)
add_bullet("Read-heavy: catalogue, books, recommendation queries.")
add_bullet("Write-heavy: cart updates during campaigns, order and payment during checkout peaks.")
add_body(
    "Scaling guidance: scale catalog/book independently from payment/shipping. "
    "Add caching for hot read endpoints."
)

add_heading("16.2 Latency Hotspots", 2)
add_bullet("Gateway fan-out views that query many services simultaneously.")
add_bullet("cart-service validation call to book-service on each cart-item insertion.")
add_body("Mitigation:")
add_bullet("Add timeout, retry for idempotent reads, and bounded concurrency in gateway.")
add_bullet("Introduce service-level metrics (p95 latency, error rate).")

add_heading("16.3 Data Growth", 2)
add_body(
    "High-growth entities: CartItem, CommentRate, Recommendation records as model "
    "refresh frequency increases."
)
add_body("Mitigation: add retention and archival policy; partitioning strategy for analytics-heavy tables.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 17. ROADMAP
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("17. Roadmap — Iteration Plan", 1)

roadmap = [
    (
        "Iteration 1 — Contract & Security Hardening",
        [
            "Add JWT authentication and role-based access checks.",
            "Publish OpenAPI (Swagger) docs per service.",
            "Add request validation constraints and standardised error schema.",
            "Implement PUT /cart-items/{id}/ for cart item quantity update.",
        ]
    ),
    (
        "Iteration 2 — Reliability Upgrade",
        [
            "Introduce retry/circuit-breaker library (e.g., tenacity).",
            "Add centralised logging (ELK stack) and distributed tracing (Jaeger/OpenTelemetry).",
            "Implement /health and /ready endpoints on all services.",
        ]
    ),
    (
        "Iteration 3 — Event-driven Order Workflow",
        [
            "Introduce RabbitMQ or Kafka message broker.",
            "Emit and consume domain events: OrderCreated, PaymentSucceeded, ShipmentCreated.",
            "Add idempotency keys to prevent duplicate payment/shipment operations.",
        ]
    ),
    (
        "Iteration 4 — Productisation",
        [
            "CI/CD pipeline with service-level automated tests.",
            "Container image scanning and dependency vulnerability checks.",
            "SLO dashboards (Grafana/Prometheus) and alerting policy.",
        ]
    ),
]
for title, items in roadmap:
    add_heading(title, 2)
    for item in items:
        add_bullet(item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 18. DISCUSSION & CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("18. Discussion & Conclusion", 1)

add_heading("18.1 Benefits Achieved", 2)
add_bullet("Clear modular decomposition aligned with DDD bounded contexts.")
add_bullet("Service-level autonomy and independent data ownership.")
add_bullet("Better extensibility for future features (auth, events, analytics).")
add_bullet("Foundation for team-level parallel development.")

add_heading("18.2 Trade-offs", 2)
add_bullet("Increased operational complexity compared to monolith.")
add_bullet("More network calls introduce potential latency.")
add_bullet("Stronger observability and resilience patterns required for production-readiness.")

add_heading("18.3 Academic vs Production Gap", 2)
add_body(
    "This implementation intentionally prioritises architecture structure and assignment-required "
    "flows. For production, additional non-functional controls are mandatory: SLOs, security "
    "hardening, async workflows, and incident tooling."
)

add_heading("18.4 Conclusion", 2)
add_body(
    "The BookStore system has been successfully decomposed into 12 microservices aligned with "
    "assignment objectives and DDD boundaries. The architecture demonstrates independent service "
    "ownership, database-per-service isolation, and REST-based collaboration with API Gateway "
    "as the controlled entry point."
)
add_body(
    "Required business flows are implemented for customer registration with cart auto-creation, "
    "cart-book validation, and rating flow. Order-payment-shipping orchestration is represented "
    "at the architecture level and can be incrementally upgraded to event-driven execution."
)
add_body(
    "Overall, the project provides a valid academic microservice foundation with a clear roadmap "
    "toward production-grade reliability, security, and observability."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# APPENDIX A — SERVICE LIST
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Appendix A — Complete Service List", 1)

app_a_data = [
    ("1",  "api-gateway",            "8000", "Entry point; proxy and UI rendering"),
    ("2",  "customer-service",       "8001", "Customer registration and management"),
    ("3",  "staff-service",          "8002", "Staff records and roles"),
    ("4",  "manager-service",        "8003", "Manager records"),
    ("5",  "catalog-service",        "8004", "Book catalogue/categories"),
    ("6",  "book-service",           "8005", "Book master data"),
    ("7",  "cart-service",           "8006", "Shopping cart and cart items"),
    ("8",  "order-service",          "8007", "Order lifecycle management"),
    ("9",  "pay-service",            "8008", "Payment record handling"),
    ("10", "ship-service",           "8009", "Shipment record handling"),
    ("11", "comment-rate-service",   "8010", "Customer reviews and ratings"),
    ("12", "recommender-ai-service", "8011", "AI-based book recommendations"),
]
tbl = doc.add_table(rows=1 + len(app_a_data), cols=4)
tbl.style = "Table Grid"
for ci, h in enumerate(["#", "Service Name", "Port", "Responsibility"]):
    shade_cell(tbl.rows[0].cells[ci], HEAD_BG)
    p = tbl.rows[0].cells[ci].paragraphs[0]; p.clear()
    r = p.add_run(h); set_font(r, bold=True, color=WHITE, size=10)
for ri, (num, name, port, resp) in enumerate(app_a_data, 1):
    cells = tbl.rows[ri].cells
    if ri % 2 == 0:
        for c in cells: shade_cell(c, ALT_BG)
    for ci, txt in enumerate([num, name, port, resp]):
        p = cells[ci].paragraphs[0]; p.clear()
        r = p.add_run(txt)
        if ci == 1:
            set_font(r, size=10, bold=True, color=TEAL, name="Courier New")
        elif ci == 2:
            set_font(r, size=10, name="Courier New")
        else:
            set_font(r, size=10)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# APPENDIX B — HTTP STATUS CODES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Appendix B — Common HTTP Status Codes", 1)

http_codes = [
    ("200 OK",                  "Successful GET — resource(s) returned"),
    ("201 Created",             "Successful POST — new resource created"),
    ("400 Bad Request",         "Serializer validation error or missing required fields"),
    ("404 Not Found",           "Resource not found (e.g., invalid book_id in cart-item)"),
    ("500 Internal Server Error","Unexpected runtime error in the service"),
    ("502 Bad Gateway",         "Gateway could not reach a downstream service"),
    ("503 Service Unavailable", "Service is temporarily down or overloaded"),
]
tbl = doc.add_table(rows=1 + len(http_codes), cols=2)
tbl.style = "Table Grid"
for ci, h in enumerate(["HTTP Status", "Meaning in this System"]):
    shade_cell(tbl.rows[0].cells[ci], HEAD_BG)
    p = tbl.rows[0].cells[ci].paragraphs[0]; p.clear()
    r = p.add_run(h); set_font(r, bold=True, color=WHITE, size=10)
for ri, (code, meaning) in enumerate(http_codes, 1):
    cells = tbl.rows[ri].cells
    if ri % 2 == 0:
        shade_cell(cells[0], ALT_BG); shade_cell(cells[1], ALT_BG)
    for ci, txt in enumerate([code, meaning]):
        p = cells[ci].paragraphs[0]; p.clear()
        r = p.add_run(txt)
        if ci == 0:
            set_font(r, size=10, bold=True, color=TEAL, name="Courier New")
        else:
            set_font(r, size=10)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# ARCHITECTURE DIAGRAMS  (picture placeholders)
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading("Architecture Diagrams", 1)
add_body(
    "The diagrams below illustrate the internal structure and inter-service relationships "
    "for each component in the BookStore microservice system. Import the corresponding "
    "architecture image into each placeholder."
)

diagram_services = [
    ("1", "API Gateway Architecture",
     "Shows gateway routing, proxy rules, and downstream service connections."),
    ("2", "Customer Service Architecture",
     "Shows Customer model, REST endpoints, and the cart-service side-effect call."),
    ("3", "Book Service Architecture",
     "Shows Book model, REST endpoints, and data flow to/from catalog-service."),
    ("4", "Cart Service Architecture",
     "Shows Cart and CartItem models, book validation call, and GET cart-detail flow."),
    ("5", "Order Service Architecture",
     "Shows Order model, status lifecycle, and relationship to pay/ship services."),
    ("6", "Payment Service Architecture",
     "Shows Payment model, order reference, and status transitions."),
    ("7", "Shipment Service Architecture",
     "Shows Shipment model, address handling, and status transitions."),
    ("8", "Comment & Rate Service Architecture",
     "Shows CommentRate model, customer/book references, and rating flow."),
    ("9", "Staff Service Architecture",
     "Shows Staff model, role field, and book management endpoints."),
    ("10", "Manager Service Architecture",
     "Shows Manager model and management endpoints."),
    ("11", "Recommender AI Service Architecture",
     "Shows Recommendation model, score computation pipeline, and gateway exposure."),
]

def add_picture_placeholder(number: str, title: str, caption: str):
    """Bordered single-cell table acting as a picture placeholder."""
    # Sub-heading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f"{number}. {title}")
    set_font(r, size=12, bold=True, color=NAVY)

    # Description line
    dp = doc.add_paragraph()
    dp.paragraph_format.space_after = Pt(6)
    dr = dp.add_run(caption)
    set_font(dr, size=10, italic=True, color=GRAY)

    # Placeholder box
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, "F5F8FC")

    # Fix cell height via XML
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "3600")   # ~6.35 cm
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

    ph_p = cell.paragraphs[0]
    ph_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Vertical space before text
    ph_p.paragraph_format.space_before = Pt(28)
    ph_p.paragraph_format.space_after  = Pt(4)

    icon_run = ph_p.add_run("[ INSERT DIAGRAM IMAGE ]")
    set_font(icon_run, size=13, bold=True, color=RGBColor(0xB0, 0xC4, 0xDE))

    sub_p = cell.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(4)
    sub_r = sub_p.add_run(f"Diagram: {title}")
    set_font(sub_r, size=9, italic=True, color=RGBColor(0x99, 0xAA, 0xBB))

    doc.add_paragraph()   # breathing room after box

for num, title, caption in diagram_services:
    add_picture_placeholder(num, title, caption)
    # Page break after every 2 diagrams to avoid crowding
    if num in ("2", "4", "6", "8", "10"):
        doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
output_path = r"l:\bookstore\Assignment05_B22DCVT527_PhamManhThang_v2.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
